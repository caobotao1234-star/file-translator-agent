# translator/scan_writer.py
# =============================================================
# 📘 教学笔记：扫描件 PDF → Word 写入器（v7.1 — 精确布局还原）
# =============================================================
# 📘 v7.1 改进：
#   - 每个单元格独立控制四条边框（per-cell borders）
#   - 每行文字独立对齐（per-line alignment）
#   - 图片按 image_position 放在文字前/后/中间
#   - 竖版文字支持（vertical text direction）
#   - 精确列宽比例
# =============================================================

import io
import os
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from core.logger import get_logger
from translator.format_engine import FormatEngine

logger = get_logger("scan_writer")

# 📘 页面可用宽度（A4 纸宽 21cm - 左右边距各 1.5cm = 18cm）
PAGE_CONTENT_WIDTH_CM = 18.0


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """
    📘 设置单元格边框
    每个参数是 dict: {"sz": 4, "val": "single", "color": "000000"}
    或 None 表示无边框
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_xml = f'<w:tcBorders {nsdecls("w")}>'
    for edge, style in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        if style:
            borders_xml += (
                f'<w:{edge} w:val="{style.get("val", "single")}" '
                f'w:sz="{style.get("sz", 4)}" w:space="0" '
                f'w:color="{style.get("color", "000000")}"/>'
            )
        else:
            borders_xml += f'<w:{edge} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    borders_xml += '</w:tcBorders>'
    tcBorders = parse_xml(borders_xml)
    tcPr.append(tcBorders)


def _set_cell_borders_from_dict(cell, borders: dict):
    """
    📘 v7.1 新增：从 Vision LLM 返回的 borders dict 设置边框
    borders: {"top": true/false, "bottom": true/false, "left": true/false, "right": true/false}
    """
    border_on = {"sz": 4, "val": "single", "color": "000000"}
    _set_cell_borders(
        cell,
        top=border_on if borders.get("top", True) else None,
        bottom=border_on if borders.get("bottom", True) else None,
        left=border_on if borders.get("left", True) else None,
        right=border_on if borders.get("right", True) else None,
    )


def _set_cell_vertical_text(cell):
    """
    📘 v7.1 新增：设置单元格文字方向为竖版（从上到下）

    📘 教学笔记：
    Word 中竖版文字通过 tcPr 的 textDirection 属性实现。
    btLr = Bottom to Top, Left to Right（竖版，从下到上读）
    tbRl = Top to Bottom, Right to Left（竖版，从上到下读，东亚竖排）
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    text_dir = parse_xml(f'<w:textDirection {nsdecls("w")} w:val="tbRlV"/>')
    tcPr.append(text_dir)


def _set_cell_width(cell, width_cm: float):
    """📘 设置单元格宽度"""
    width_emu = int(width_cm * 360000)  # 1cm = 360000 EMU
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_emu}" w:type="dxa"/>')
    tcPr.append(tcW)


def _get_alignment(align_str: str):
    """📘 字符串 → WD_ALIGN_PARAGRAPH 枚举"""
    mapping = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }
    return mapping.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)


def _get_font_size(size_str: str) -> float:
    """📘 字号描述 → pt 值"""
    mapping = {"small": 8, "normal": 10.5, "large": 14, "title": 18}
    return mapping.get(size_str, 10.5)


def _set_run_font(run, bold: bool = False, size_pt: float = 10, font_name: str = "Microsoft YaHei",
                  color_hex: str = None):
    """
    📘 设置 run 的字体属性
    color_hex: 可选，如 "#FF0000" 表示红色文字
    """
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color_hex and color_hex.startswith("#") and len(color_hex) == 7:
        try:
            r = int(color_hex[1:3], 16)
            g = int(color_hex[3:5], 16)
            b = int(color_hex[5:7], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        except (ValueError, IndexError):
            pass


def _add_image_to_cell(cell, image_bytes: bytes, align: str = "center", max_width_cm: float = None):
    """
    📘 v7.2 改进：把裁剪好的图片嵌入到表格单元格中
    max_width_cm 可由调用方指定，默认 4cm
    """
    try:
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(image_bytes))
        w_px, h_px = img.size
        max_w = max_width_cm or 4.0
        w_cm = min(max_w, w_px * 2.54 / 200)

        para = cell.add_paragraph()
        para.alignment = _get_alignment(align)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run()
        img_stream = io.BytesIO(image_bytes)
        run.add_picture(img_stream, width=Cm(w_cm))
        return para
    except Exception as e:
        logger.warning(f"单元格图片嵌入失败: {e}")
        return None


def _add_table_to_doc(doc: Document, table_data: dict, translations: Dict[str, str],
                      page_idx: int, elem_idx: int):
    """
    📘 教学笔记：把结构化表格数据写入 Word 文档（v7.2 完整版）

    📘 v7.2 改进：
    1. 支持表格级 borders（true/false），作为所有单元格的默认值
    2. 单元格级 borders 可覆盖表格级设置
    3. 无边框表格 = 布局网格（"一切布局皆表格"方法论的核心）
    4. 每行文字独立对齐（per-line alignment via "lines" array）
    5. 图片按 image_position 放在文字前/后
    6. 竖版文字支持（vertical text direction）
    7. 精确列宽比例
    """
    rows_data = table_data.get("rows", [])
    if not rows_data:
        return

    col_widths_pct = table_data.get("col_widths", [])

    # 📘 表格级 borders 默认值（Brain 可能在 table 级别设置 borders: true/false）
    table_borders_default = table_data.get("borders", True)

    # 📘 计算实际列数（考虑 colspan）
    max_cols = 0
    for row in rows_data:
        cells = row.get("cells", row) if isinstance(row, dict) else row
        if isinstance(cells, dict):
            cells = cells.get("cells", [])
        col_count = sum(cell.get("colspan", 1) for cell in cells)
        max_cols = max(max_cols, col_count)

    num_rows = len(rows_data)
    if max_cols == 0 or num_rows == 0:
        return

    # 📘 计算列宽（cm）
    if col_widths_pct and len(col_widths_pct) == max_cols:
        total_pct = sum(col_widths_pct)
        col_widths_cm = [PAGE_CONTENT_WIDTH_CM * p / total_pct for p in col_widths_pct]
    else:
        col_widths_cm = [PAGE_CONTENT_WIDTH_CM / max_cols] * max_cols

    # 创建表格
    table = doc.add_table(rows=num_rows, cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 📘 设置表格总宽度
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    total_width_twips = int(PAGE_CONTENT_WIDTH_CM * 567)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_width_twips}" w:type="dxa"/>')
    tblPr.append(tblW)

    # 📘 v7.2: 无边框表格 = 布局网格，减小单元格内边距使布局更紧凑
    if table_borders_default is False:
        cell_margin_xml = (
            f'<w:tblCellMar {nsdecls("w")}>'
            f'<w:top w:w="0" w:type="dxa"/>'
            f'<w:left w:w="28" w:type="dxa"/>'
            f'<w:bottom w:w="0" w:type="dxa"/>'
            f'<w:right w:w="28" w:type="dxa"/>'
            f'</w:tblCellMar>'
        )
        tblPr.append(parse_xml(cell_margin_xml))

    # 📘 跟踪被 rowspan 占用的位置
    occupied = [[False] * max_cols for _ in range(num_rows)]

    for row_idx, row in enumerate(rows_data):
        cells = row.get("cells", row) if isinstance(row, dict) else row
        if isinstance(cells, dict):
            cells = cells.get("cells", [])

        col_cursor = 0
        cell_data_idx = 0

        for cell_data in cells:
            while col_cursor < max_cols and occupied[row_idx][col_cursor]:
                col_cursor += 1
            if col_cursor >= max_cols:
                break

            colspan = cell_data.get("colspan", 1)
            rowspan = cell_data.get("rowspan", 1)
            cell_bold = cell_data.get("bold", False)
            cell_align = cell_data.get("align", "left")
            is_vertical = cell_data.get("vertical", False)

            # 📘 获取文字内容：支持 "lines" 数组（per-line alignment）或 "text" 简写
            cell_lines = cell_data.get("lines")
            if cell_lines and isinstance(cell_lines, list):
                original_text = "\n".join(l.get("text", "") for l in cell_lines)
            else:
                original_text = cell_data.get("text", "").strip()
                cell_lines = None

            # 📘 查找译文
            key = f"pg{page_idx}_e{elem_idx}_r{row_idx}_c{cell_data_idx}"
            translated = translations.get(key, original_text)

            # 获取单元格
            cell = table.cell(row_idx, col_cursor)

            # 📘 合并单元格
            if colspan > 1 or rowspan > 1:
                end_row = min(row_idx + rowspan - 1, num_rows - 1)
                end_col = min(col_cursor + colspan - 1, max_cols - 1)
                try:
                    cell = cell.merge(table.cell(end_row, end_col))
                except Exception as e:
                    logger.debug(f"合并失败 [{row_idx},{col_cursor}]->[{end_row},{end_col}]: {e}")

            # 📘 竖版文字
            if is_vertical:
                _set_cell_vertical_text(cell)

            # 📘 图片处理
            has_image = cell_data.get("has_image", False)
            cropped_image = cell_data.get("cropped_image")
            image_position = cell_data.get("image_position", "after")

            # 📘 写入内容
            cell.text = ""

            # 📘 图片在文字前面
            if has_image and cropped_image and image_position == "before":
                _add_image_to_cell(cell, cropped_image, cell_align)

            # 📘 写入文字（支持 per-line alignment）
            trans_lines = translated.split("\n")
            # 📘 v7.3: 单元格文字颜色支持（detect_colors 工具检测到的颜色）
            font_color = cell_data.get("font_color")

            if cell_lines and len(cell_lines) == len(trans_lines):
                # 📘 per-line alignment：每行用原文的对齐方式
                for line_idx, (trans_text, orig_line) in enumerate(zip(trans_lines, cell_lines)):
                    line_align = orig_line.get("align", cell_align)
                    if line_idx == 0:
                        para = cell.paragraphs[0]
                    else:
                        para = cell.add_paragraph()
                    para.alignment = _get_alignment(line_align)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(1)
                    run = para.add_run(trans_text.strip())
                    _set_run_font(run, bold=cell_bold, size_pt=10, color_hex=font_color)
            else:
                # 📘 简写模式：所有行用同一个对齐
                for line_idx, line_text in enumerate(trans_lines):
                    if line_idx == 0:
                        para = cell.paragraphs[0]
                    else:
                        para = cell.add_paragraph()
                    para.alignment = _get_alignment(cell_align)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(1)
                    run = para.add_run(line_text.strip())
                    _set_run_font(run, bold=cell_bold, size_pt=10, color_hex=font_color)

            # 📘 图片在文字后面（默认）
            if has_image and cropped_image and image_position != "before":
                _add_image_to_cell(cell, cropped_image, cell_align)

            # 📘 设置列宽（仅第一行设置即可）
            if row_idx == 0 and col_cursor < len(col_widths_cm):
                w = sum(col_widths_cm[col_cursor:col_cursor + colspan])
                _set_cell_width(cell, w)

            # 📘 v7.2: 边框处理（3 层优先级）
            # 1. 单元格级 borders dict → 最高优先级，per-cell 精确控制
            # 2. 单元格级 borders bool → 覆盖表格默认
            # 3. 表格级 table_borders_default → 全局默认
            # 📘 "一切布局皆表格"方法论的关键：borders: false = 无边框布局网格
            cell_borders = cell_data.get("borders")
            if isinstance(cell_borders, dict):
                _set_cell_borders_from_dict(cell, cell_borders)
            elif cell_borders is False or (cell_borders is None and table_borders_default is False):
                # 📘 无边框 = 布局网格
                _set_cell_borders(cell, top=None, bottom=None, left=None, right=None)
            else:
                # 📘 有边框（cell_borders=True 或 table_borders_default=True）
                border_on = {"sz": 4, "val": "single", "color": "000000"}
                _set_cell_borders(cell, top=border_on, bottom=border_on,
                                  left=border_on, right=border_on)

            # 标记被占用的位置
            for r in range(row_idx, min(row_idx + rowspan, num_rows)):
                for c in range(col_cursor, min(col_cursor + colspan, max_cols)):
                    if r != row_idx or c != col_cursor:
                        occupied[r][c] = True

            col_cursor += colspan
            cell_data_idx += 1

    # 📘 v7.2: 有边框表格后加间距，无边框布局表格不加（避免布局间隙）
    if table_borders_default is not False:
        doc.add_paragraph()


def _add_paragraph_to_doc(doc: Document, elem: dict, translations: Dict[str, str],
                          page_idx: int, elem_idx: int):
    """📘 把段落文本写入 Word 文档（带格式）"""
    key = f"pg{page_idx}_e{elem_idx}_para"
    text = translations.get(key, elem.get("text", "").strip())

    # 📘 教学笔记：段落中的 \n 拆分为多个 Word 段落
    # Brain 有时仍会用 \n 分隔多行内容。拆分后每行独立段落，
    # 避免 Word 中出现一个超长段落。
    lines = text.split("\n") if "\n" in text else [text]
    align = _get_alignment(elem.get("align", "left"))
    size_pt = _get_font_size(elem.get("font_size", "normal"))
    is_bold = elem.get("bold", False)

    # 📘 v7.3: 段落文字颜色支持
    font_color = elem.get("font_color")

    for line_text in lines:
        line_text = line_text.strip()
        if not line_text:
            continue
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(line_text)
        _set_run_font(run, bold=is_bold, size_pt=size_pt, color_hex=font_color)


def _add_signature_block_to_doc(doc: Document, elem: dict, translations: Dict[str, str],
                                 page_idx: int, elem_idx: int):
    """
    📘 教学笔记：签名/落款区域写入

    签名区域通常是居中或右对齐的多行文字块，
    如医生签名、公司落款、日期等。
    Brain 输出 {"type": "signature_block", "align": "center", "lines": [...]}
    """
    align = _get_alignment(elem.get("align", "center"))
    sig_lines = elem.get("lines", [])

    # 📘 签名区域前加一个空行
    doc.add_paragraph()

    for line_data in sig_lines:
        if isinstance(line_data, str):
            line_text = line_data
            line_bold = False
        else:
            line_text = line_data.get("text", "").strip()
            line_bold = line_data.get("bold", False)

        if not line_text:
            continue

        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(line_text)
        _set_run_font(run, bold=line_bold, size_pt=10)


def _add_spacer_to_doc(doc: Document):
    """📘 添加空行间距"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)


def _add_image_to_doc(doc: Document, elem: dict):
    """
    📘 教学笔记：嵌入裁剪的图片到 Word 文档

    📘 v7.2 改进：支持 crop_key 引用 + bbox_pct 自动裁剪
    Brain 通过 crop_image_region 工具裁剪后，图片 bytes 在后处理阶段
    被嵌入到 elem["cropped_image"] 中。

    图片宽度根据 bbox_pct 的宽度比例自适应：
    - 小图（< 30% 页宽）→ 限制 5cm
    - 中图（30-60%）→ 限制 10cm
    - 大图（> 60%）→ 限制 16cm
    """
    cropped = elem.get("cropped_image")
    description = elem.get("description", "图片区域")
    bbox_pct = elem.get("bbox_pct")

    if cropped:
        para = doc.add_paragraph()
        # 📘 根据 bbox_pct 推断对齐方式
        align = "center"
        if bbox_pct and len(bbox_pct) == 4:
            center_x = (bbox_pct[0] + bbox_pct[2]) / 2
            if center_x < 35:
                align = "left"
            elif center_x > 65:
                align = "right"
        para.alignment = _get_alignment(align)

        run = para.add_run()
        img_stream = io.BytesIO(cropped)
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(cropped))
            w_px, h_px = img.size

            # 📘 根据 bbox_pct 宽度比例决定 Word 中的图片宽度
            if bbox_pct and len(bbox_pct) == 4:
                region_width_pct = bbox_pct[2] - bbox_pct[0]
                max_w_cm = PAGE_CONTENT_WIDTH_CM * region_width_pct / 100
                max_w_cm = max(2.0, min(max_w_cm, 16.0))  # 至少 2cm，最多 16cm
            else:
                max_w_cm = 8.0

            w_cm = min(max_w_cm, w_px * 2.54 / 200)
            run.add_picture(img_stream, width=Cm(w_cm))
        except Exception:
            run.add_picture(img_stream, width=Cm(6))
    else:
        # 📘 没有裁剪图片，用文字占位
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[{description}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True


def write_scan_pdf(
    parsed_data: Dict[str, Any],
    translations: Dict[str, str],
    output_path: str,
    format_engine: FormatEngine,
    source_path: str = None,
    layout_overrides: Dict[str, dict] = None,
    scan_mode: str = "adaptive",
    fixed_fontsize: float = 10.0,
):
    """
    📘 扫描件写入主函数（v7.1 — 精确布局还原）

    生成 Word 文档，每页包含：
    1. 原始页面图片（参考）
    2. 结构化译文（per-cell 边框、per-line 对齐、图片位置、竖版文字）
    3. 分页符

    📘 教学笔记：翻译查找策略
    translations 的 key 由 Brain 生成（如 pg0_e1_r0_c0），
    但 Brain 的 key 命名经常与 writer 的 key 构造不一致。
    所以除了 key 匹配，还建立 {原文: 译文} 的文本匹配作为 fallback。
    """
    page_structures = parsed_data.get("page_structures", [])
    page_images = parsed_data.get("page_images", [])

    if not page_structures:
        raise ValueError("parsed_data 中缺少 page_structures")

    # 📘 输出路径强制 .docx
    base, ext = os.path.splitext(output_path)
    if ext.lower() != ".docx":
        output_path = base + ".docx"

    logger.info(f"开始生成扫描件 Word 文档: {output_path}")
    print(f"[📝 扫描件写入] 生成 Word 文档（精确布局）...", flush=True)

    doc = Document()

    # 📘 默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # 📘 窄边距
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    translated_count = 0
    embedded_count = 0

    for page_idx, structure in enumerate(page_structures):
        if page_idx > 0:
            doc.add_page_break()

        # 📘 页面标题
        heading = doc.add_heading(f"第 {page_idx + 1} 页", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 📘 嵌入原始页面图片（参考）
        if page_idx < len(page_images) and page_images[page_idx]:
            img_stream = io.BytesIO(page_images[page_idx])
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_stream, width=Cm(16))

            sep = doc.add_paragraph()
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sep.add_run("— 以下为译文 —")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

        # 📘 按结构写入内容
        elements = structure.get("elements", [])
        for elem_idx, elem in enumerate(elements):
            elem_type = elem.get("type", "")

            if elem_type == "table":
                _add_table_to_doc(doc, elem, translations, page_idx, elem_idx)
                # 📘 教学笔记：翻译计数（key 匹配 + 嵌入匹配）
                # key 匹配：Brain 的 key 与 writer 构造的 key 一致
                # 嵌入匹配：key 不一致但 scan_agent 已把译文嵌入 elem["text"]
                for row_idx, row in enumerate(elem.get("rows", [])):
                    cells = row.get("cells", row) if isinstance(row, dict) else row
                    if isinstance(cells, dict):
                        cells = cells.get("cells", [])
                    for col_idx, cell in enumerate(cells):
                        key = f"pg{page_idx}_e{elem_idx}_r{row_idx}_c{col_idx}"
                        if key in translations:
                            translated_count += 1
                        else:
                            # 📘 嵌入的译文也算翻译成功
                            cell_text = cell.get("text", "").strip()
                            if cell_text:
                                embedded_count += 1

            elif elem_type == "paragraph":
                _add_paragraph_to_doc(doc, elem, translations, page_idx, elem_idx)
                key = f"pg{page_idx}_e{elem_idx}_para"
                if key in translations:
                    translated_count += 1
                else:
                    para_text = elem.get("text", "").strip()
                    if para_text:
                        embedded_count += 1

            elif elem_type == "signature_block":
                # 📘 签名/落款区域
                _add_signature_block_to_doc(doc, elem, translations, page_idx, elem_idx)
                for line_data in elem.get("lines", []):
                    line_text = line_data.get("text", "") if isinstance(line_data, dict) else str(line_data)
                    if line_text.strip():
                        embedded_count += 1

            elif elem_type == "spacer":
                # 📘 空行间距
                _add_spacer_to_doc(doc)

            elif elem_type == "image_region":
                _add_image_to_doc(doc, elem)

    doc.save(output_path)

    total_written = translated_count + embedded_count
    logger.info(
        f"扫描件 Word 文档生成完成: {output_path} "
        f"(key匹配 {translated_count} + 嵌入匹配 {embedded_count} = 共 {total_written} 个单元)"
    )
    print(
        f"[✅ 扫描件写入完成] 生成 Word 文档，"
        f"翻译了 {total_written} 个单元 (key:{translated_count} + 嵌入:{embedded_count})",
        flush=True,
    )

    return output_path


def write_overlay_pdf(
    overlay_images: Dict[int, bytes],
    page_images: list,
    output_path: str,
    num_pages: int,
):
    """
    📘 教学笔记：保留背景模式 — 图片合成 PDF

    把处理后的页面图片（已覆盖译文）合成为 PDF 文件。
    如果某页没有被 overlay 处理，使用原始页面图片。

    📘 技术方案：
    - 用 Pillow 把每页 JPEG 转为 PDF 页面
    - 第一页用 save()，后续页用 append_images
    - 输出路径强制 .pdf（保留背景 = 保留原始视觉效果）
    """
    from PIL import Image as PILImage

    base, ext = os.path.splitext(output_path)
    output_path = base + ".pdf"

    logger.info(f"开始生成保留背景 PDF: {output_path}")
    print(f"[📝 保留背景] 合成 PDF（{num_pages} 页）...", flush=True)

    pdf_pages = []
    for page_idx in range(num_pages):
        # 📘 优先使用 overlay 处理后的图片，否则用原图
        if page_idx in overlay_images:
            img_bytes = overlay_images[page_idx]
        elif page_idx < len(page_images) and page_images[page_idx]:
            img_bytes = page_images[page_idx]
        else:
            logger.warning(f"第 {page_idx} 页无图片数据，跳过")
            continue

        pil_img = PILImage.open(io.BytesIO(img_bytes)).convert("RGB")
        pdf_pages.append(pil_img)

    if not pdf_pages:
        raise ValueError("没有可用的页面图片")

    # 📘 Pillow save as PDF：第一页 save，其余 append
    pdf_pages[0].save(
        output_path,
        "PDF",
        resolution=200.0,
        save_all=True,
        append_images=pdf_pages[1:] if len(pdf_pages) > 1 else [],
    )

    logger.info(f"保留背景 PDF 生成完成: {output_path} ({len(pdf_pages)} 页)")
    print(
        f"[✅ 保留背景完成] 生成 PDF，{len(pdf_pages)} 页",
        flush=True,
    )

    return output_path
