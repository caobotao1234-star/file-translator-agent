# translator/docx_writer.py
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from typing import List, Dict, Any, Optional
from translator.format_engine import FormatEngine
from core.logger import get_logger

# =============================================================
# 📘 教学笔记：Word 文档生成器（v4 — 段落 + 表格统一写回）
# =============================================================
# v3 只处理段落，表格里的文字不翻译。
#
# v4 的策略：
#   - 段落：和 v3 一样，支持 Run 级别标记
#   - 表格：通过 key（如 "t_0_1_2"）定位到具体单元格的段落
#   - 统一的写回逻辑：不管是段落还是表格单元格，
#     都是"找到目标段落 → 替换 Run 文本 → 映射字体"
# =============================================================

logger = get_logger("docx_writer")

RUN_TAG_PATTERN = re.compile(r'<r(\d+)>(.*?)</r\1>', re.DOTALL)


def _parse_tagged_text(text: str) -> Optional[Dict[int, str]]:
    """解析带标记的译文，返回 {Run编号: 译文} 字典。"""
    matches = RUN_TAG_PATTERN.findall(text)
    if not matches:
        return None
    result = {}
    for idx_str, content in matches:
        result[int(idx_str)] = content
    return result


def _remap_run_font(run, format_engine: FormatEngine, style_name: str = None):
    """只替换 Run 的字体名，其他格式全部保留。"""
    original_font = run.font.name
    if original_font:
        new_font = format_engine.resolve_font(original_font, style_name)
        if new_font:
            run.font.name = new_font
            if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), new_font)


def _replace_paragraph_text(paragraph, translated_text: str, is_tagged: bool,
                            format_engine: FormatEngine, style_name: str = None):
    """
    替换一个段落的文本（通用逻辑，段落和表格单元格共用）。
    """
    original_runs = paragraph.runs

    if len(original_runs) == 0:
        clean = RUN_TAG_PATTERN.sub(r'\2', translated_text)
        paragraph.add_run(clean)
        return

    # ---- 多 Run + 带标记：逐 Run 替换 ----
    if is_tagged and len(original_runs) > 1:
        run_texts = _parse_tagged_text(translated_text)

        if run_texts is not None:
            non_empty_indices = [ri for ri, run in enumerate(original_runs) if run.text]

            success = True
            for tag_idx, text in run_texts.items():
                if tag_idx < len(non_empty_indices):
                    actual_idx = non_empty_indices[tag_idx]
                    original_runs[actual_idx].text = text
                    _remap_run_font(original_runs[actual_idx], format_engine, style_name)
                else:
                    logger.warning(f"标记 r{tag_idx} 超出 Run 范围，降级处理")
                    success = False
                    break

            if success:
                tagged_indices = set()
                for tag_idx in run_texts.keys():
                    if tag_idx < len(non_empty_indices):
                        tagged_indices.add(non_empty_indices[tag_idx])
                for ri, run in enumerate(original_runs):
                    if run.text and ri not in tagged_indices:
                        run.text = ""
                return

        # 标记解析失败，降级
        logger.warning("标记解析失败，降级为整段替换")

    # ---- 单 Run 或降级：整段替换 ----
    clean_text = RUN_TAG_PATTERN.sub(r'\2', translated_text)
    original_runs[0].text = clean_text
    _remap_run_font(original_runs[0], format_engine, style_name)
    for run in original_runs[1:]:
        run.text = ""


def _resolve_paragraph_from_key(doc, key: str):
    """
    根据 key 找到对应的段落对象。

    key 格式：
      "p_3"           → doc.paragraphs[3]
      "t_0_1_2"       → doc.tables[0].rows[1].cells[2].paragraphs[0]
      "t_0_1_2_1"     → doc.tables[0].rows[1].cells[2].paragraphs[1]
    """
    parts = key.split("_")

    if parts[0] == "p":
        para_idx = int(parts[1])
        if para_idx < len(doc.paragraphs):
            return doc.paragraphs[para_idx]
        return None

    if parts[0] == "t":
        table_idx = int(parts[1])
        row_idx = int(parts[2])
        col_idx = int(parts[3])
        para_idx = int(parts[4]) if len(parts) > 4 else 0

        if table_idx >= len(doc.tables):
            return None
        table = doc.tables[table_idx]
        if row_idx >= len(table.rows):
            return None
        row = table.rows[row_idx]
        if col_idx >= len(row.cells):
            return None
        cell = row.cells[col_idx]
        if para_idx >= len(cell.paragraphs):
            return None
        return cell.paragraphs[para_idx]

    return None


def write_docx(
    parsed_data: Dict[str, Any],
    translations: Dict[str, str],
    output_path: str,
    format_engine: FormatEngine,
    source_path: str = None,
):
    """
    基于原文档生成翻译后的 Word 文档（段落 + 表格统一处理）。

    参数：
        parsed_data: docx_parser.parse_docx() 的返回值
        translations: {key: 翻译后文本} 字典
        output_path: 输出文件路径
        format_engine: 格式映射引擎
        source_path: 原文档路径
    """
    if not source_path:
        raise ValueError("source_path 不能为空，需要原文档来克隆格式")

    logger.info(f"开始生成文档: {output_path}")
    doc = Document(source_path)

    replaced_count = 0
    for item in parsed_data["items"]:
        key = item["key"]

        if item.get("is_empty") or key not in translations:
            continue

        paragraph = _resolve_paragraph_from_key(doc, key)
        if paragraph is None:
            logger.warning(f"无法定位 key={key}，跳过")
            continue

        translated_text = translations[key]
        is_tagged = item.get("tagged_text", False)
        style_name = item.get("style", {}).get("style_name")

        _replace_paragraph_text(paragraph, translated_text, is_tagged,
                                format_engine, style_name)
        replaced_count += 1

    doc.save(output_path)
    logger.info(f"文档生成完成: {output_path}（替换了 {replaced_count} 个翻译单元）")
