# translator/docx_parser.py
from docx import Document
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from typing import List, Dict, Any, Optional
from core.logger import get_logger

# =============================================================
# 📘 教学笔记：Word 文档解析器（v3 — 段落 + 表格统一处理）
# =============================================================
# v2 只处理 doc.paragraphs，表格里的文字完全跳过。
#
# v3 的策略：
#   - 按文档的实际顺序遍历所有"块级元素"（段落 + 表格）
#   - 段落：和 v2 一样，支持 Run 级别标记
#   - 表格：遍历每个单元格，每个单元格内部也是段落，同样支持标记
#   - 用统一的 key 体系标识每个翻译单元：
#       "p_3"       → 第3个段落
#       "t_0_1_2"   → 第0个表格、第1行、第2列
#       "t_0_1_2_1" → 第0个表格、第1行、第2列、第1个段落（单元格内多段落时）
#
# Word 文档的 XML 结构：
#   <w:body>
#     <w:p>段落1</w:p>
#     <w:tbl>表格1</w:tbl>    ← 不在 doc.paragraphs 里！
#     <w:p>段落2</w:p>
#   </w:body>
#
#   我们通过遍历 doc.element.body 的子元素来按顺序处理。
# =============================================================

logger = get_logger("docx_parser")


def _extract_run_format(run) -> Dict[str, Any]:
    """提取一个 Run 的格式信息"""
    fmt = {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "font_name": run.font.name,
        "font_size": run.font.size.pt if run.font.size else None,
        "font_color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
    }
    return fmt


def _extract_paragraph_style(paragraph) -> Dict[str, Any]:
    """提取段落级别的样式信息"""
    style = {
        "style_name": paragraph.style.name if paragraph.style else None,
        "alignment": str(paragraph.alignment) if paragraph.alignment else None,
    }
    pf = paragraph.paragraph_format
    if pf:
        style["line_spacing"] = pf.line_spacing
        style["space_before"] = pf.space_before.pt if pf.space_before else None
        style["space_after"] = pf.space_after.pt if pf.space_after else None
        style["first_line_indent"] = pf.first_line_indent.pt if pf.first_line_indent else None
    return style


def _build_tagged_text(runs_data: List[Dict]) -> str:
    """把多个 Run 的文本用标记包裹"""
    parts = []
    for i, run in enumerate(runs_data):
        parts.append(f"<r{i}>{run['text']}</r{i}>")
    return "".join(parts)


def _parse_paragraph(para, key: str) -> Optional[Dict[str, Any]]:
    """
    解析一个段落（可以是文档级段落，也可以是表格单元格内的段落）。
    返回 None 表示空段落。
    """
    raw_text = para.text.strip()

    if not raw_text:
        return None

    runs = []
    for run in para.runs:
        if not run.text:
            continue
        runs.append({
            "text": run.text,
            "format": _extract_run_format(run),
        })

    # 判断是否需要 Run 标记
    needs_tagging = False
    if len(runs) > 1:
        first_fmt = runs[0]["format"]
        for r in runs[1:]:
            if r["format"] != first_fmt:
                needs_tagging = True
                break

    full_text = _build_tagged_text(runs) if needs_tagging else raw_text

    return {
        "key": key,
        "type": "paragraph",
        "full_text": full_text,
        "style": _extract_paragraph_style(para),
        "runs": runs,
        "tagged_text": needs_tagging,
    }


def _parse_table(table, table_idx: int) -> List[Dict[str, Any]]:
    """
    解析一个表格，返回所有需要翻译的单元格段落列表。

    每个单元格可能包含多个段落，我们逐个处理。
    key 格式：t_{表格序号}_{行号}_{列号} 或 t_{表格序号}_{行号}_{列号}_{段落号}
    """
    items = []
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_paragraphs = cell.paragraphs
            if len(cell_paragraphs) == 1:
                # 单段落单元格（最常见）
                key = f"t_{table_idx}_{row_idx}_{col_idx}"
                result = _parse_paragraph(cell_paragraphs[0], key)
                if result:
                    result["type"] = "table_cell"
                    items.append(result)
            else:
                # 多段落单元格
                for para_idx, para in enumerate(cell_paragraphs):
                    key = f"t_{table_idx}_{row_idx}_{col_idx}_{para_idx}"
                    result = _parse_paragraph(para, key)
                    if result:
                        result["type"] = "table_cell"
                        items.append(result)
    return items


def parse_docx(filepath: str) -> Dict[str, Any]:
    """
    解析 Word 文档，按文档顺序返回所有翻译单元（段落 + 表格单元格）。

    返回格式：
    {
        "items": [
            {"key": "p_0", "type": "paragraph", "full_text": "...", ...},
            {"key": "t_0_1_2", "type": "table_cell", "full_text": "...", ...},
            ...
        ]
    }
    """
    logger.info(f"开始解析文档: {filepath}")
    doc = Document(filepath)

    items = []
    para_idx = 0
    table_idx = 0

    # 按文档 XML 顺序遍历所有块级元素
    for element in doc.element.body:
        tag = element.tag.split('}')[-1]  # 去掉命名空间前缀

        if tag == 'p':
            # 段落元素
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                key = f"p_{para_idx}"
                result = _parse_paragraph(para, key)
                if result:
                    items.append(result)
                else:
                    # 空段落也记录（保持结构信息）
                    items.append({
                        "key": key,
                        "type": "paragraph",
                        "full_text": "",
                        "style": _extract_paragraph_style(para),
                        "runs": [],
                        "tagged_text": False,
                        "is_empty": True,
                    })
            para_idx += 1

        elif tag == 'tbl':
            # 表格元素
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                table_items = _parse_table(table, table_idx)
                items.extend(table_items)
            table_idx += 1

    # 统计
    para_count = sum(1 for i in items if i["type"] == "paragraph" and not i.get("is_empty"))
    cell_count = sum(1 for i in items if i["type"] == "table_cell")
    tagged_count = sum(1 for i in items if i.get("tagged_text"))

    logger.info(f"解析完成: {para_count} 个段落，{cell_count} 个表格单元格，"
                f"{tagged_count} 个含格式标记")

    return {"items": items}
